"""
Word 渲染器(Sprint 6 - v0.8.0)

职责:
- 使用 docxtpl 渲染模板变量(保留 Word 原生样式)
- 使用 python-docx 在"补充条款"占位点插入 AI 补充条款段落
- 导出 .docx 到 uploads/generated/{uuid}.docx
- 返回 (file_path, file_name, file_size)

技术:
- docxtpl:基于 Jinja2 的 Word 模板渲染,保留字体/表格/页眉页脚样式
- python-docx:已安装(Sprint 4 知识库 DOCX 加载依赖),用于段落插入

渲染策略:
1. DocxTemplate(template_path) 加载模板
2. 构建 context:user_vars(用户变量)+ 补充条款拼接文本
3. template.render(context) 渲染占位符
4. 若存在 {{supplementary_clauses}} 占位符,直接渲染为条款文本;
   否则在文档末尾追加"补充条款"段落(AI 生成的条款)
5. 保存到 uploads/generated/{uuid}.docx

约束:
- 不修改模板原文,每次生成基于模板副本渲染
- 文件 UUID 命名,存 uploads/generated/
- 禁止 print() / return str(e)
"""
import os
import uuid
from datetime import datetime
from flask import current_app

from app.extensions.logger import logger
from app.utils.exceptions import BusinessError


# ---------- 配置常量 ----------
_GENERATED_SUBDIR = 'generated'  # 生成文件子目录(相对 UPLOAD_FOLDER)
_SUPPLEMENTARY_CLAUSES_VAR = 'supplementary_clauses'  # 补充条款占位符变量名


def _get_generated_dir():
    """获取生成文件目录(uploads/generated/),并确保目录存在。"""
    gen_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], _GENERATED_SUBDIR)
    os.makedirs(gen_dir, exist_ok=True)
    return gen_dir


def _build_clauses_text(generated_clauses):
    """
    将 AI 补充条款拼接为 Word 可渲染的文本

    格式:
        第一条 付款条款
        <条款正文>

        第二条 违约责任
        <条款正文>
        ...

    :param generated_clauses: [{name, content, source, references}]
    :return: str 拼接文本(无条款返回空串)
    """
    if not generated_clauses:
        return ''
    parts = []
    cn_nums = ['第一条', '第二条', '第三条', '第四条', '第五条',
               '第六条', '第七条', '第八条', '第九条', '第十条',
               '第十一条', '第十二条', '第十三条', '第十四条', '第十五条']
    for i, clause in enumerate(generated_clauses):
        name = clause.get('name', f'条款{i + 1}')
        content = clause.get('content', '').strip()
        if not content:
            continue
        prefix = cn_nums[i] if i < len(cn_nums) else f'第{i + 1}条'
        parts.append(f'{prefix} {name}\n{content}')
    return '\n\n'.join(parts)


def render_contract(template_path, input_variables, generated_clauses,
                    output_title=None):
    """
    渲染合同 Word 文档

    :param template_path: 模板文件路径
    :param input_variables: 用户填写的变量键值 {var_name: value}
    :param generated_clauses: AI 补充条款 [{name, content, source, references}]
    :param output_title: 输出文件标题(可选,用于生成文件名)
    :return: dict {file_path, file_name, file_size, clause_count}
    :raises BusinessError: 渲染失败
    """
    if not template_path or not os.path.exists(template_path):
        raise BusinessError('模板文件不存在,无法渲染')

    try:
        from docxtpl import DocxTemplate
        from docx import Document
    except ImportError:
        logger.exception('[Gen:renderer] docxtpl / python-docx 未安装')
        raise BusinessError('Word 渲染依赖未安装(docxtpl / python-docx)')

    # ---------- 1. 加载模板 ----------
    try:
        doc = DocxTemplate(template_path)
    except Exception as e:
        logger.exception('[Gen:renderer] 模板加载失败: %s', template_path)
        raise BusinessError(f'模板文件加载失败: {e}')

    # ---------- 2. 构建渲染上下文 ----------
    # 用户变量(字符串化,避免 docxtpl 处理非字符串类型异常)
    context = {}
    for k, v in (input_variables or {}).items():
        context[k] = '' if v is None else str(v)

    # 补充条款文本
    clauses_text = _build_clauses_text(generated_clauses or [])
    # 若模板含 supplementary_clauses 占位符,直接填充;否则置空(后续追加段落)
    context[_SUPPLEMENTARY_CLAUSES_VAR] = clauses_text

    clause_count = len([c for c in (generated_clauses or []) if c.get('content', '').strip()])

    # ---------- 3. 渲染占位符 ----------
    try:
        doc.render(context)
    except Exception as e:
        logger.exception('[Gen:renderer] 模板渲染失败')
        raise BusinessError(f'模板渲染失败: {e}')

    # ---------- 4. 追加 AI 补充条款段落(若模板无 supplementary_clauses 占位符) ----------
    # 策略:检测渲染后文档是否仍含未替换的 supplementary_clauses 文本;
    #       若模板无此占位符且存在 AI 条款,则在文档末尾追加"补充条款"段落
    if clauses_text and not _has_supplementary_var(template_path):
        try:
            _append_clauses_section(doc, generated_clauses)
        except Exception:
            logger.exception('[Gen:renderer] 追加补充条款段落失败(忽略,继续保存)')

    # ---------- 5. 保存到 uploads/generated/{uuid}.docx ----------
    gen_dir = _get_generated_dir()
    saved_filename = f'{uuid.uuid4().hex}.docx'
    output_path = os.path.join(gen_dir, saved_filename)

    try:
        doc.save(output_path)
        file_size = os.path.getsize(output_path)
    except Exception as e:
        logger.exception('[Gen:renderer] 文档保存失败: %s', output_path)
        # 清理可能产生的残留文件
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except OSError:
                pass
        raise BusinessError(f'文档保存失败: {e}')

    logger.info('[Gen:renderer] Word 渲染成功: output=%s size=%s clauses=%s',
                output_path, file_size, clause_count)

    return {
        'file_path': output_path,
        'file_name': saved_filename,
        'file_size': file_size,
        'clause_count': clause_count,
    }


def _has_supplementary_var(template_path):
    """
    检测模板是否含 {{supplementary_clauses}} 占位符
    :return: bool
    """
    try:
        from docxtpl import DocxTemplate
        doc = DocxTemplate(template_path)
        var_names = doc.get_undeclared_template_variables()
        return _SUPPLEMENTARY_CLAUSES_VAR in var_names
    except Exception:
        return False


def _append_clauses_section(doc, generated_clauses):
    """
    在文档末尾追加"补充条款"段落(AI 生成的条款)

    使用 python-docx 操作底层文档(doc.docx 为 python-docx Document 对象)
    - 追加分页符前的空行
    - 追加"补充条款"标题
    - 逐条追加条款(条款名加粗 + 正文)

    :param doc: DocxTemplate 实例
    :param generated_clauses: [{name, content, source, references}]
    """
    # DocxTemplate.docx 为底层 python-docx Document 对象
    document = doc.docx if hasattr(doc, 'docx') else None
    if document is None:
        return

    cn_nums = ['第一条', '第二条', '第三条', '第四条', '第五条',
               '第六条', '第七条', '第八条', '第九条', '第十条',
               '第十一条', '第十二条', '第十三条', '第十四条', '第十五条']

    # 追加空行
    document.add_paragraph('')

    # 追加"补充条款"标题(加粗)
    heading = document.add_heading('补充条款', level=2)

    idx = 0
    for clause in generated_clauses or []:
        name = clause.get('name', '')
        content = clause.get('content', '').strip()
        if not content:
            continue
        idx += 1
        prefix = cn_nums[idx - 1] if idx - 1 < len(cn_nums) else f'第{idx}条'
        # 条款名(加粗段落)
        title_para = document.add_paragraph()
        title_run = title_para.add_run(f'{prefix} {name}')
        title_run.bold = True
        # 条款正文
        document.add_paragraph(content)
