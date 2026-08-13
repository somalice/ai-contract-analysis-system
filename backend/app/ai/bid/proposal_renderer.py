"""
投标 Word 渲染器(Sprint 7 - v0.9.0)

职责:
- 复用 docxtpl + python-docx 渲染投标 Word 文档
- 封面占位符(project_name / tender_org / bidder_company / budget / deadline)用 docxtpl 渲染
- 各章节(technical / commercial / responsive / qualification / summary)用 python-docx 追加
- 模板 templates/proposal_base.docx 首次运行时懒生成(避免提交二进制文件)
- 导出 .docx 到 uploads/proposals/{uuid}.docx

技术:
- docxtpl:基于 Jinja2 的 Word 模板渲染(复用 Sprint 6 已安装依赖)
- python-docx:已安装(Sprint 4 知识库 DOCX 加载依赖),用于段落插入

渲染策略:
1. 确保 proposal_base.docx 存在(不存在则懒生成)
2. DocxTemplate(template_path) 加载模板
3. 构建 context:封面占位符
4. template.render(context) 渲染封面
5. 追加各章节(章节名加粗 + 正文 Markdown 转换为段落)
6. 保存到 uploads/proposals/{uuid}.docx

约束:
- 不修改模板原文,每次生成基于模板副本渲染
- 文件 UUID 命名,存 uploads/proposals/
- 禁止 print() / return str(e)
- 不修改 Sprint 6 word_renderer.py
"""
import os
import uuid
from datetime import datetime

from flask import current_app

from app.extensions.logger import logger
from app.utils.exceptions import BusinessError


# ---------- 配置常量 ----------
_PROPOSAL_SUBDIR = 'proposals'  # 投标生成文件子目录(相对 UPLOAD_FOLDER)

# ---------- 封面占位符变量 ----------
_COVER_VARS = (
    'project_name', 'tender_org', 'bidder_company', 'budget', 'deadline',
)

# ---------- 章节中文名 ----------
_SECTION_NAMES = {
    'technical': '技术方案',
    'commercial': '商务文件',
    'responsive': '响应文件',
    'qualification': '资质文件',
    'summary': '投标摘要',
}


def _get_proposal_dir():
    """获取投标生成文件目录(uploads/proposals/),并确保目录存在。"""
    gen_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], _PROPOSAL_SUBDIR)
    os.makedirs(gen_dir, exist_ok=True)
    return gen_dir


def _get_template_path():
    """
    获取投标模板路径(templates/proposal_base.docx)
    若模板不存在,则懒生成(用 python-docx 写封面占位符文本)
    """
    # 模板路径:app/templates/proposal_base.docx(与 Flask template_folder 一致)
    template_dir = os.path.join(current_app.root_path, 'templates')
    template_path = os.path.join(template_dir, 'proposal_base.docx')

    if os.path.exists(template_path):
        return template_path

    # ---------- 懒生成模板 ----------
    logger.info('[Bid:renderer] 投标模板不存在,懒生成: %s', template_path)
    try:
        _generate_base_template(template_path)
    except Exception:
        logger.exception('[Bid:renderer] 模板懒生成失败: %s', template_path)
        raise BusinessError('投标模板生成失败,请联系管理员')
    return template_path


def _generate_base_template(template_path: str):
    """
    生成基础投标模板(封面 + 章节占位符)
    - 标题:投标文件
    - 封面占位符:{{project_name}} / {{tender_org}} / {{bidder_company}} / {{budget}} / {{deadline}}
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 设置默认字体(中文)
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # ---------- 封面 ----------
    # 主标题
    title = doc.add_heading('投标文件', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 封面占位符段落
    doc.add_paragraph('')  # 空行
    for var in _COVER_VARS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('{{' + var + '}}')
        run.font.size = Pt(14)
        run.bold = True

    # 分页
    doc.add_page_break()

    # ---------- 章节占位符提示 ----------
    note = doc.add_paragraph()
    note_run = note.add_run('（以下章节由系统自动生成）')
    note_run.italic = True
    note_run.font.size = Pt(10)

    doc.save(template_path)
    logger.info('[Bid:renderer] 投标模板已生成: %s', template_path)


def _render_markdown_to_docx(document, content: str):
    """
    将 Markdown 文本渲染为 python-docx 段落
    - ## 二级标题 → add_heading(level=2)
    - ### 三级标题 → add_heading(level=3)
    - - 列表项 → add_paragraph(style='List Bullet')
    - **加粗** → run.bold = True
    - 其他 → add_paragraph

    :param document: python-docx Document 对象
    :param content: Markdown 文本
    """
    if not content:
        return

    lines = content.split('\n')
    for line in lines:
        line = line.rstrip()
        if not line:
            # 空行:跳过(避免过多空段落)
            continue

        # 二级标题
        if line.startswith('## '):
            heading_text = line[3:].strip()
            document.add_heading(heading_text, level=2)
            continue

        # 三级标题
        if line.startswith('### '):
            heading_text = line[4:].strip()
            document.add_heading(heading_text, level=3)
            continue

        # 四级标题
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            document.add_heading(heading_text, level=4)
            continue

        # 列表项(- 或 *)
        if line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            _add_paragraph_with_bold(document, list_text, style='List Bullet')
            continue

        # 有序列表项(1. 2. 等)
        import re
        ordered_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if ordered_match:
            list_text = ordered_match.group(2).strip()
            _add_paragraph_with_bold(document, list_text, style='List Number')
            continue

        # 普通段落(处理 **加粗**)
        _add_paragraph_with_bold(document, line)


def _add_paragraph_with_bold(document, text: str, style: str = None):
    """
    添加段落,处理 **加粗** 标记
    :param document: python-docx Document
    :param text: 文本(可能含 **加粗** 标记)
    :param style: 段落样式(可选)
    """
    if not text:
        return
    # 分割 ** 加粗标记
    import re
    parts = re.split(r'\*\*(.+?)\*\*', text)
    if len(parts) == 1:
        # 无加粗标记
        if style:
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(text)
        return

    # 有加粗标记:逐段添加 run
    p = document.add_paragraph(style=style) if style else document.add_paragraph()
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        # 奇数索引(即 ** 之间的内容)加粗
        if i % 2 == 1:
            run.bold = True


def render_proposal(bid_info: dict, requirements: dict,
                    company_profile: dict, generated_sections: list,
                    output_title: str = None) -> dict:
    """
    渲染投标 Word 文档

    :param bid_info: 招标文件信息(供封面 project_name / tender_org)
    :param requirements: 招标需求(供封面 budget / deadline)
    :param company_profile: 企业资料(供封面 bidder_company)
    :param generated_sections: AI 生成章节 [{section_type, section_name, content, source, references}]
    :param output_title: 输出文件标题(可选,用于生成文件名)
    :return: dict {file_path, file_name, file_size, section_count}
    :raises BusinessError: 渲染失败
    """
    # ---------- 0. 确保模板存在 ----------
    template_path = _get_template_path()

    try:
        from docxtpl import DocxTemplate
        from docx import Document
    except ImportError:
        logger.exception('[Bid:renderer] docxtpl / python-docx 未安装')
        raise BusinessError('Word 渲染依赖未安装(docxtpl / python-docx)')

    # ---------- 1. 加载模板 ----------
    try:
        doc = DocxTemplate(template_path)
    except Exception as e:
        logger.exception('[Bid:renderer] 模板加载失败: %s', template_path)
        raise BusinessError(f'模板文件加载失败: {e}')

    # ---------- 2. 构建封面渲染上下文 ----------
    context = {}
    req_data = requirements or {}
    bid_data = bid_info or {}
    profile_data = company_profile or {}

    # 封面占位符(从需求 / 招标信息 / 企业资料 中提取)
    context['project_name'] = str(req_data.get('project_name') or bid_data.get('title') or '（项目名称）')
    context['tender_org'] = str(req_data.get('tender_org') or '（招标单位）')
    context['bidder_company'] = str(profile_data.get('company_name') or '（投标单位）')
    context['budget'] = str(req_data.get('budget') or '（预算金额）')
    context['deadline'] = str(req_data.get('deadline') or '（截止时间）')

    # ---------- 3. 渲染封面占位符 ----------
    try:
        doc.render(context)
    except Exception as e:
        logger.exception('[Bid:renderer] 封面渲染失败')
        raise BusinessError(f'封面渲染失败: {e}')

    # ---------- 4. 追加各章节(用 python-docx) ----------
    document = doc.docx if hasattr(doc, 'docx') else None
    if document is None:
        raise BusinessError('无法访问 Word 文档对象(python-docx)')

    # 按 sort_order 排序章节
    section_order = {'technical': 1, 'commercial': 2, 'responsive': 3,
                     'qualification': 4, 'summary': 5}
    sorted_sections = sorted(
        generated_sections or [],
        key=lambda s: section_order.get(s.get('section_type'), 99)
    )

    section_count = 0
    for section in sorted_sections:
        section_name = section.get('section_name', '')
        content = section.get('content', '').strip()
        if not content:
            continue

        section_count += 1
        # 分页(每个章节新起一页,除第一章外)
        if section_count > 1:
            document.add_page_break()
        # 章节标题(一级标题)
        document.add_heading(section_name, level=1)
        # 章节正文(Markdown 渲染)
        _render_markdown_to_docx(document, content)

    # ---------- 5. 保存到 uploads/proposals/{uuid}.docx ----------
    gen_dir = _get_proposal_dir()
    saved_filename = f'{uuid.uuid4().hex}.docx'
    output_path = os.path.join(gen_dir, saved_filename)

    try:
        doc.save(output_path)
        file_size = os.path.getsize(output_path)
    except Exception as e:
        logger.exception('[Bid:renderer] 文档保存失败: %s', output_path)
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except OSError:
                pass
        raise BusinessError(f'文档保存失败: {e}')

    logger.info('[Bid:renderer] Word 渲染成功: output=%s size=%s sections=%s',
                output_path, file_size, section_count)

    return {
        'file_path': output_path,
        'file_name': saved_filename,
        'file_size': file_size,
        'section_count': section_count,
    }
