"""
DOCX Loader(Sprint 4 - v0.6.0)

职责:
- 用 python-docx 提取 .docx 文档文本
- docx 无明确页码,视为单页(page_number=1);按段落提取并保留段落分隔

设计说明:
- 提取所有段落文本,空段落转为段落分隔(双换行),便于后续 chunk 按段落切分
- 表格文本也提取(按行拼接),避免知识丢失
"""
from docx import Document

from app.utils.text_utils import clean_text
from app.extensions.logger import logger
from .base import BaseLoader, Page


class DocxLoader(BaseLoader):
    """DOCX 文档加载器(python-docx)"""

    @property
    def supported_extensions(self) -> tuple:
        return ('docx',)

    def load(self, file_path: str) -> list:
        parts = []
        try:
            doc = Document(file_path)

            # 1. 段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
                else:
                    # 空段落作为段落分隔
                    parts.append('')

            # 2. 表格(按行拼接单元格)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))

            logger.info('[Knowledge:docx_loader] DOCX 解析完成: %s 段落/行', len(parts))
        except Exception as e:
            logger.exception('[Knowledge:docx_loader] DOCX 解析失败: %s', file_path)
            raise

        # 合并为单页文本(段落间双换行);clean_text 会规范空行
        raw = '\n\n'.join(parts)
        text = clean_text(raw)
        return [Page(page_number=1, text=text)]
